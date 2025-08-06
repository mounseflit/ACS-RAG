import pypdf
import docx
import pandas as pd
import io
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
# from ibm_watson import SpeechToTextV1
from ibm_cloud_sdk_core.authenticators import IAMAuthenticator
import os

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, "rb") as f:
        reader = pypdf.PdfReader(f)
        for page_num in range(len(reader.pages)):
            text += reader.pages[page_num].extract_text() or ""
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

def extract_text_from_xlsx(file_path):
    xls = pd.ExcelFile(file_path)
    text = []
    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name)
        text.append(f"Sheet: {sheet_name}\n")
        text.append(df.to_string())
    return "\n".join(text)

def extract_text_from_image(image_path):
    try:
        return pytesseract.image_to_string(Image.open(image_path))
    except Exception as e:
        return f"Error during OCR: {e}"

def process_scanned_pdf(file_path):
    text = ""
    images = convert_from_path(file_path)
    for i, image in enumerate(images):
        text += f"\n--- Page {i+1} ---\n"
        text += pytesseract.image_to_string(image)
    return text

# def transcribe_audio(audio_path, api_key, service_url):
#     try:
#         authenticator = IAMAuthenticator(api_key)
#         speech_to_text = SpeechToTextV1(
#             authenticator=authenticator
#         )
#         speech_to_text.set_service_url(service_url)

#         with open(audio_path, 'rb') as audio_file:
#             speech_recognition_results = speech_to_text.recognize(
#                 audio=audio_file,
#                 content_type='audio/wav',
#                 model='en-US_BroadbandModel',
#                 continuous=True
#             ).get_result()
        
#         text = ""
#         for result in speech_recognition_results["results"]:
#             text += result["alternatives"][0]["transcript"]
#         return text

#     except Exception as e:
#         return f"Error during audio transcription: {e}"

def chunk_text(text, chunk_size=500, overlap=50):
    chunks = []
    words = text.split()
    i = 0
    while i < len(words):
        chunk = " ".join(words[i:i + chunk_size])
        chunks.append(chunk)
        i += chunk_size - overlap
    return chunks

def process_file(file_path, file_type, watson_stt_api_key=None, watson_stt_service_url=None):
    content = ""
    if file_type == "pdf":
        text = extract_text_from_pdf(file_path)
        if len(text.strip()) < 100: # Arbitrary threshold for sparse text, indicating scanned PDF
            content = process_scanned_pdf(file_path)
        else:
            content = text
    elif file_type == "docx":
        content = extract_text_from_docx(file_path)
    elif file_type == "xlsx":
        content = extract_text_from_xlsx(file_path)
    elif file_type == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
    elif file_type == "image":
        content = extract_text_from_image(file_path)
    # elif file_type == "audio":
    #     if watson_stt_api_key and watson_stt_service_url:
    #         content = transcribe_audio(file_path, watson_stt_api_key, watson_stt_service_url)
    #     else:
    #         content = "Audio processing requires Watson Speech to Text API key and service URL."
    else:
        return [] # Return empty list for unsupported types

    if content:
        chunks = chunk_text(content)
        processed_data = []
        for i, chunk in enumerate(chunks):
            processed_data.append({
                "content": chunk,
                "metadata": {
                    "file_name": os.path.basename(file_path),
                    "file_type": file_type,
                    "chunk_id": i
                }
            })
        return processed_data
    return []


# if __name__ == "__main__":
#     # Example Usage (create dummy files for testing)
#     with open("test.txt", "w") as f:
#         f.write("This is a test text file. This file contains multiple sentences to demonstrate chunking. Let's see how it works with a longer text.")
    
#     doc = docx.Document()
#     doc.add_paragraph("This is a test docx file. It also has some content for chunking.")
#     doc.save("test.docx")

#     df = pd.DataFrame({
#         'col1': [1, 2, 3, 4, 5],
#         'col2': [6, 7, 8, 9, 10]
#     })
#     df.to_excel("test.xlsx", index=False)

#     # Create a dummy PDF file (requires reportlab)
#     from reportlab.pdfgen import canvas
#     c = canvas.Canvas("test.pdf")
#     c.drawString(100, 750, "This is a test PDF file. This is the first sentence. This is the second sentence. This is the third sentence.")
#     c.save()

#     # Create a dummy image file for OCR testing
#     from PIL import Image, ImageDraw, ImageFont
#     img = Image.new('RGB', (200, 50), color = (255, 255, 255))
#     d = ImageDraw.Draw(img)
#     d.text((10,10), "Hello OCR! This is a test image for OCR.", fill=(0,0,0))
#     img.save("test_ocr.png")

#     # Create a dummy audio file (requires pydub and audiogen, but we'll just create a dummy file for now)
#     # For actual testing, a real .wav file would be needed.
#     with open("dummy_audio.wav", "w") as f:
#         f.write("This is a dummy audio file content.")

#     print("\n--- Processing test.txt ---")
#     print(process_file("test.txt", "txt"))

#     print("\n--- Processing test.docx ---")
#     print(process_file("test.docx", "docx"))

#     print("\n--- Processing test.xlsx ---")
#     print(process_file("test.xlsx", "xlsx"))

#     print("\n--- Processing test.pdf ---")
#     print(process_file("test.pdf", "pdf"))

#     print("\n--- Processing test_ocr.png ---")
#     print(process_file("test_ocr.png", "image"))

#     print("\n--- Processing dummy_audio.wav (placeholder) ---")
#     # To test audio, replace None with actual API key and URL
#     print(process_file("dummy_audio.wav", "audio", watson_stt_api_key=None, watson_stt_service_url=None))



